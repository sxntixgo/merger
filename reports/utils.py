from django.core.files import File

from .models import ReportTemplate
from main.models import Attach, Proj, Vuln, RISK

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import matplotlib
matplotlib.use('agg')
from matplotlib import pyplot
from matplotlib.colors import CSS4_COLORS

from datetime import date
from re import match
from os import path, remove
from wsgiref.util import FileWrapper
import random

RISK_COLORS = [ 'risk_info_color','risk_low_color', 'risk_medium_color', 'risk_high_color', 'risk_critical_color']
SECTIONS = {
    'sec_es': 'Executive Summary',
    'sec_sor': 'Summary of Results',
    'sec_method': 'Methodology',
    'sec_find': 'Findings',
    'sec_conc': 'Conclusion',
    'sec_desc': 'Description',
    'sec_evid': 'Evidence',
    'sec_sol': 'Solution',
}
STYLES = ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Normal', 'Caption']


def color_to_rgb(font_color):
    r_hex_str = CSS4_COLORS[font_color][1:3]
    g_hex_str = CSS4_COLORS[font_color][3:5]
    b_hex_str = CSS4_COLORS[font_color][5:7]
    
    return RGBColor(int(r_hex_str, base=16), int(g_hex_str, base=16), int(b_hex_str, base=16))

# Remove bottom border
# Source: https://github.com/python-openxml/python-docx/issues/105
def remove_border(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:sz'), '0')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'white')
    pBdr.append(bottom)

def get_donut(path, proj, template):
    values = []
    labels = []
    colors = []

    for color in RISK_COLORS:
        colors.append(template.COLORS[getattr(template, color)][1])

    for risk in RISK:
        values.append(Vuln.objects.filter(risk=risk[0], proj=proj.pk).count())
        labels.append(risk[1])

    # Reverse the data so that criticals are printed first.
    colors.reverse()
    values.reverse()
    labels.reverse()

    fig, ax = pyplot.subplots(figsize=(8, 4), subplot_kw=dict(aspect='equal'))
    wedges, texts, autotexts = ax.pie(values, colors=colors, wedgeprops=dict(width=0.5, edgecolor='w'), autopct=lambda p: '{:.1f}%'.format(round(p)) if p > 0 else '', pctdistance=1.3)
    ax.legend(wedges, labels, title="Risk", loc="center right", bbox_to_anchor=(1.25, 0, 0.5, 1))

    pyplot.setp(autotexts, size=14, weight="bold")
    pyplot.savefig(path + 'donut.png')

# Modified from the source
# --- from https://github.com/python-openxml/python-docx/issues/590,
# --- mods by CD
def iter_heading(paragraphs):
    for paragraph in paragraphs:
        isItHeading=match('Heading ([1-9])', paragraph.style.name)
        if isItHeading:
            yield int(isItHeading.groups()[0]), paragraph

def addHeaderNumbering(document):
    hNums = [0] * 10 # To include the 9 possible headings
    for index,hx in iter_heading(document.paragraphs):
        # ---put zeroes below---
        for i in range(index + 1, len(hNums)):
            hNums[i] = 0
        # ---increment this---
        hNums[index] += 1
        # ---prepare the string---
        hStr = ''

        for i in range(1 , index + 1):
            hStr += '{}.'.format(hNums[i])
        # ---add the numbering---
        hx.text = hStr+ " " + hx.text

def generate_document(request, path):
    proj_name = request.POST.get('proj_name')
    proj = Proj.objects.filter(name=proj_name).first()
    template_name = request.POST.get('report_template_name')
    template = ReportTemplate.objects.filter(name=template_name).first()

    document = Document()

    for style in STYLES:
        font = document.styles[style].font
        style = style.lower().replace(' ', '_')
        font.name = getattr(template, style + '_font')
        font.size = Pt(getattr(template, style + '_size'))
        font.bold = getattr(template, style + '_bold')
        font.italic = getattr(template, style + '_italic')
        font_color = template.COLORS[getattr(template, style + '_color')][1]
        font.color.rgb = color_to_rgb(font_color)

    cover = template.sec_cover
    if cover == 0:
        run = document.add_paragraph().add_run()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        for element in ['cover_title', 'cover_company_name']:
            text = getattr(template, element)
            run = document.add_paragraph().add_run(text)
            run.add_break()
            run.add_break()
            font = run.font
            font.name = getattr(template, 'title_font')
            font.size = Pt(36)
            font_color = template.COLORS[getattr(template, 'title_color')][1]
            font.color.rgb = color_to_rgb(font_color)

        for element in ['cover_contact_name', 'cover_contact_email', 'cover_contact_phone_number']:
            text = getattr(template, element)
            run = document.add_paragraph().add_run(str(text))
            font = run.font
            font.name = getattr(template, 'normal_font')
            font.size = Pt(12)
            font_color = template.COLORS[getattr(template, 'normal_color')][1]
            font.color.rgb = color_to_rgb(font_color)

        run = document.add_paragraph().add_run()
        run.add_break()
        run.add_break()

        text = date.today().strftime("%b-%d-%Y")
        par = document.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(text)
        font = run.font
        font.name = getattr(template, 'normal_font')
        font.size = Pt(12)
        font_color = template.COLORS[getattr(template, 'normal_color')][1]
        font.color.rgb = color_to_rgb(font_color)
        
        run.add_break(WD_BREAK.PAGE) # page break
        
        if template.sec_toc:
            p = document.add_heading('Table of Contents', 1)
            
            # To add ToC
            # From https://stackoverflow.com/questions/18595864/python-create-a-table-of-contents-with-python-docx-lxml
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')  # creates a new element
            fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:t')
            fldChar3.text = "Right-click to update field."
            fldChar2.append(fldChar3)

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(fldChar4)
            p_element = paragraph._p

            run.add_break(WD_BREAK.PAGE) # page break

    elif cover == 1:
        title = template.cover_title
        p = document.add_heading(title, 0)
        remove_border(p)

    for heading in SECTIONS.keys():
        if hasattr(template, heading):
            paragraph = document.add_heading(SECTIONS[heading], 1)
            if heading == 'sec_es':
                text = getattr(template, 'es_text')
                if text:
                    document.add_paragraph(text)
            if heading == 'sec_sor':
                # Add donut
                get_donut(path, proj, template)
                document.add_picture(path + 'donut.png', width=Inches(6))
                remove(path + 'donut.png')

                # Add table header
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Title'
                hdr_cells[1].text = 'Risk (Score)'
                hdr_cells[2].text = 'Description'
                hdr_cells[3].text = 'Solution'
                
                # Add table findings
                vulns = Vuln.objects.filter(proj=proj.pk)
                for vuln in vulns:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vuln.title
                    row_cells[1].text = '{} ({})'.format(RISK[vuln.risk][1], vuln.score)
                    row_cells[2].text = vuln.description
                    row_cells[3].text = vuln.solution

            if heading == 'sec_method':
                text = getattr(template, 'method_text')
                if text:
                    document.add_paragraph(text)

            if heading == 'sec_find':
                vulns = Vuln.objects.filter(proj=proj.pk)
                for vuln in vulns:             
                    document.add_heading(('{} [{}]').format(vuln.title, RISK[vuln.risk][1]), 2)
                    document.add_paragraph('{} {}'.format('CVE:', vuln.score), 'List Bullet')
                    document.add_paragraph('{} {}'.format('Score:', vuln.score), 'List Bullet')
                    document.add_paragraph(vuln.description)
                    for subheading in ['sec_evid', 'sec_sol']:
                        document.add_heading(SECTIONS[subheading], 3)
                        if subheading == 'sec_evid':
                            document.add_paragraph(vuln.evidence)
                            attachments = Attach.objects.filter(vuln=vuln).all()
                            for attachment in attachments:
                                document.add_picture(path + str(attachment.media), width=Inches(6))
                                document.add_paragraph('Figure' + ': ' + attachment.caption, style='Caption')
                        else:
                            document.add_paragraph(vuln.solution)

            if heading == 'sec_conc':
                text = getattr(template, 'conc_text')
                if text:
                    document.add_paragraph(text)

    addHeaderNumbering(document)

    hex_chars = '0123456789abcdef'
    document_name =  proj_name + '_report_' + ''.join(random.choice(hex_chars) for n in range(6)) + '.docx'
    document.save(path + document_name)

    return document_name